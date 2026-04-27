import re
from docx import Document


def read_transcript(file_path):
    with open(file_path, 'r') as file:
        return file.readlines()


def extract_meeting_info(transcript_lines):
    participants = []
    agenda_items = []
    discussed_topics = []
    agreed_commits = []

    for line in transcript_lines:
        if line.startswith('Participant: '):
            participants.append(line.replace('Participant: ', '').strip())
        elif line.startswith('Agenda: '):
            agenda_items.append(line.replace('Agenda: ', '').strip())
        elif line.startswith('Discussed: '):
            discussed_topics.append(line.replace('Discussed: ', '').strip())
        elif line.startswith('Commit: '):
            agreed_commits.append(line.replace('Commit: ', '').strip())

    return participants, agenda_items, discussed_topics, agreed_commits


def generate_document(participants, agenda_items, discussed_topics, agreed_commits, output_path):
    doc = Document()
    doc.add_heading('Meeting Minutes', level=1)
    doc.add_heading('Participants', level=2)
    for participant in participants:
        doc.add_paragraph(participant)

    doc.add_heading('Agenda Items', level=2)
    for item in agenda_items:
        doc.add_paragraph(item)

    doc.add_heading('Discussed Topics', level=2)
    for topic in discussed_topics:
        doc.add_paragraph(topic)

    doc.add_heading('Agreed Commitments', level=2)
    for commit in agreed_commits:
        doc.add_paragraph(commit)

    doc.save(output_path)


if __name__ == '__main__':
    transcript_file = 'transcript.txt'  # Path to the transcript file
    output_file = 'meeting_minutes.docx'  # Desired output DOCX file path

    transcript_lines = read_transcript(transcript_file)
    participants, agenda_items, discussed_topics, agreed_commits = extract_meeting_info(transcript_lines)
    generate_document(participants, agenda_items, discussed_topics, agreed_commits, output_file)
